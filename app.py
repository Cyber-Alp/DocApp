import io
import os
import re
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docxtpl import DocxTemplate
from openpyxl import load_workbook

app = Flask(__name__)

# Ensure the storage folder exists
TEMPLATE_DIR = "server_templates"
os.makedirs(TEMPLATE_DIR, exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

# NEW: Send the list of templates to the phone
@app.route('/api/templates', methods=['GET'])
def get_templates():
    files = os.listdir(TEMPLATE_DIR)
    docx_files = [f for f in files if f.endswith('.docx')]
    xlsx_files = [f for f in files if f.endswith('.xlsx')]
    return jsonify({"docx": docx_files, "xlsx": xlsx_files})

# NEW: Allow uploading new templates from the phone
@app.route('/api/upload_template', methods=['POST'])
def upload_template():
    file = request.files['file']
    if file:
        save_path = os.path.join(TEMPLATE_DIR, file.filename)
        file.save(save_path)
        return jsonify({"status": "success"})
    return jsonify({"status": "error"})

@app.route('/api/process', methods=['POST'])
def process():
    # NOW: We get the template name from the phone, not the file itself
    template_name = request.form['template_name']
    format_type = request.form['format']
    action = request.form['action']
    
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    
    if not os.path.exists(template_path):
        return jsonify({"status": "error", "message": "Template not found on server."})

    # Read the file from the hard drive directly into RAM
    with open(template_path, 'rb') as f:
        stream = io.BytesIO(f.read())

    try:
        if format_type == 'docx':
            if action == 'extract':
                doc = Document(stream)
                text = "\n".join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += "\n" + cell.text
                variables = list(set(re.findall(r"\{\{(.*?)\}\}", text)))
                return jsonify({"status": "success", "variables": variables})
            
            elif action == 'generate':
                context = {k: v[0] for k, v in request.form.to_dict(flat=False).items() if k not in ['format', 'action', 'template_name']}
                stream.seek(0)
                doc = DocxTemplate(stream)
                doc.render(context)
                
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                return send_file(output, as_attachment=True, download_name=f"Generated_{template_name}", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        elif format_type == 'xlsx':
            if action == 'extract':
                wb = load_workbook(stream)
                text = ""
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value: text += str(cell.value) + "\n"
                variables = list(set(re.findall(r"\{\{(.*?)\}\}", text)))
                return jsonify({"status": "success", "variables": variables})
            
            elif action == 'generate':
                context = {k: v[0] for k, v in request.form.to_dict(flat=False).items() if k not in ['format', 'action', 'template_name']}
                stream.seek(0)
                wb = load_workbook(stream)
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                for key, value in context.items():
                                    cell.value = cell.value.replace(f"{{{{{key}}}}}", str(value))
                
                output = io.BytesIO()
                wb.save(output)
                output.seek(0)
                return send_file(output, as_attachment=True, download_name=f"Generated_{template_name}", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

if __name__ == '__main__':
    print("Server running! Go to http://127.0.0.1:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)